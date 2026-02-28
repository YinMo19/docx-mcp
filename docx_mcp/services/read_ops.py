from __future__ import annotations

import re
from collections.abc import Iterator
from typing import Any

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from docx_mcp.errors import DocxMCPError
from docx_mcp.services.document_io import (
    list_docx_files,
    open_document,
    resolve_directory,
    resolve_docx_path,
)


def iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:  # pragma: no cover - internal usage only
        raise TypeError(f"Unsupported parent type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def list_available_documents(directory: str | None = None) -> dict[str, Any]:
    resolved_dir = resolve_directory(directory)
    docs = list_docx_files(resolved_dir)
    return {
        "directory": str(resolved_dir),
        "documents": [
            {
                "name": doc.name,
                "path": str(doc),
                "size_bytes": doc.stat().st_size,
            }
            for doc in docs
        ],
        "count": len(docs),
    }


def _collect_all_text_lines(doc: DocxDocument) -> list[str]:
    lines: list[str] = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        else:
            for row in block.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
    return lines


def get_document_info(filename: str) -> dict[str, Any]:
    path = resolve_docx_path(filename)
    doc = open_document(path)
    core = doc.core_properties
    lines = _collect_all_text_lines(doc)

    return {
        "path": str(path),
        "title": core.title,
        "author": core.author,
        "subject": core.subject,
        "keywords": core.keywords,
        "created": str(core.created) if core.created else None,
        "modified": str(core.modified) if core.modified else None,
        "last_modified_by": core.last_modified_by,
        "revision": core.revision,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "word_count": len(" ".join(lines).split()),
        "line_count": len(lines),
    }


def get_document_text(filename: str) -> dict[str, Any]:
    path = resolve_docx_path(filename)
    doc = open_document(path)
    lines = _collect_all_text_lines(doc)
    return {
        "path": str(path),
        "text": "\n".join(lines),
        "line_count": len(lines),
    }


def _table_preview(table: Table, preview_rows: int = 3, preview_cols: int = 3) -> list[list[str]]:
    preview: list[list[str]] = []
    for row in table.rows[:preview_rows]:
        row_values: list[str] = []
        for cell in row.cells[:preview_cols]:
            row_values.append(cell.text.strip())
        preview.append(row_values)
    return preview


def get_document_outline(filename: str) -> dict[str, Any]:
    path = resolve_docx_path(filename)
    doc = open_document(path)

    paragraphs: list[dict[str, Any]] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        paragraphs.append(
            {
                "index": idx,
                "text": text,
                "style_id": paragraph.style.style_id if paragraph.style else None,
                "style_name": paragraph.style.name if paragraph.style else None,
            }
        )

    tables: list[dict[str, Any]] = []
    for idx, table in enumerate(doc.tables):
        row_count = len(table.rows)
        col_count = max((len(row.cells) for row in table.rows), default=0)
        tables.append(
            {
                "index": idx,
                "rows": row_count,
                "columns": col_count,
                "preview": _table_preview(table),
            }
        )

    return {
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _build_pattern(text_to_find: str, whole_word: bool) -> str:
    escaped = re.escape(text_to_find)
    if whole_word:
        return rf"\b{escaped}\b"
    return escaped


def _match_in_text(
    content: str,
    pattern: re.Pattern[str],
    location: str,
    occurrences: list[dict[str, Any]],
) -> None:
    for match in pattern.finditer(content):
        start, end = match.span()
        context_start = max(0, start - 30)
        context_end = min(len(content), end + 30)
        occurrences.append(
            {
                "location": location,
                "position": start,
                "context": content[context_start:context_end],
            }
        )


def find_text_in_document(
    filename: str,
    text_to_find: str,
    match_case: bool = False,
    whole_word: bool = False,
) -> dict[str, Any]:
    if not text_to_find:
        raise DocxMCPError(
            code="INVALID_QUERY",
            message="text_to_find cannot be empty",
        )

    path = resolve_docx_path(filename)
    doc = open_document(path)

    flags = 0 if match_case else re.IGNORECASE
    regex = re.compile(_build_pattern(text_to_find, whole_word), flags)
    occurrences: list[dict[str, Any]] = []

    for p_idx, paragraph in enumerate(doc.paragraphs):
        content = paragraph.text or ""
        if content:
            _match_in_text(
                content=content,
                pattern=regex,
                location=f"Paragraph {p_idx}",
                occurrences=occurrences,
            )

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                content = cell.text or ""
                if content:
                    _match_in_text(
                        content=content,
                        pattern=regex,
                        location=f"Table {t_idx}, Row {r_idx}, Column {c_idx}",
                        occurrences=occurrences,
                    )

    return {
        "query": text_to_find,
        "match_case": match_case,
        "whole_word": whole_word,
        "occurrences": occurrences,
        "total_count": len(occurrences),
        "path": str(path),
    }

from __future__ import annotations

from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_mcp.errors import DocxMCPError
from docx_mcp.services.document_io import open_document, resolve_for_write, save_document

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _validate_alignment(value: str, field: str) -> WD_ALIGN_PARAGRAPH:
    normalized = _ALIGNMENT_MAP.get(value.lower())
    if normalized is None:
        raise DocxMCPError(
            code="INVALID_ALIGNMENT",
            message=f"Unsupported {field}: {value}",
            details={"supported": sorted(_ALIGNMENT_MAP.keys())},
        )
    return normalized


def _append_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def _set_start_page_number(section, start_page_number: int | None) -> None:
    sect_pr = section._sectPr
    current = sect_pr.find(qn("w:pgNumType"))
    if current is not None:
        sect_pr.remove(current)
    if start_page_number is None:
        return
    if start_page_number < 1:
        raise DocxMCPError(
            code="INVALID_PAGE_NUMBER",
            message="start_page_number must be >= 1.",
            details={"start_page_number": start_page_number},
        )
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(start_page_number))
    sect_pr.append(pg_num)


def _prepare_paragraph(container, *, clear_existing: bool):
    if not container.paragraphs:
        return container.add_paragraph()
    if clear_existing:
        for paragraph in container.paragraphs:
            paragraph.clear()
        return container.paragraphs[0]
    return container.paragraphs[-1]


def set_headers_footers(
    filename: str,
    *,
    section_indices: list[int] | None = None,
    header_text: str | None = None,
    footer_text: str | None = None,
    header_alignment: str = "center",
    footer_alignment: str = "center",
    include_page_number: bool = False,
    clear_existing: bool = True,
    start_page_number: int | None = None,
    different_first_page: bool | None = None,
    different_odd_even: bool | None = None,
    unlink_from_previous: bool = True,
    output_filename: str | None = None,
) -> dict[str, Any]:
    normalized_header_alignment = _validate_alignment(
        header_alignment, field="header_alignment"
    )
    normalized_footer_alignment = _validate_alignment(
        footer_alignment, field="footer_alignment"
    )

    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)

    if different_odd_even is not None:
        doc.settings.odd_and_even_pages_header_footer = different_odd_even

    selected_indices = (
        list(range(len(doc.sections))) if section_indices is None else list(section_indices)
    )
    for idx in selected_indices:
        if idx < 0 or idx >= len(doc.sections):
            raise DocxMCPError(
                code="SECTION_INDEX_OUT_OF_RANGE",
                message=f"Section index out of range: {idx}",
                details={"section_count": len(doc.sections)},
            )

    updated_indices: list[int] = []
    for idx in selected_indices:
        section = doc.sections[idx]
        if unlink_from_previous:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
        if different_first_page is not None:
            section.different_first_page_header_footer = different_first_page

        if header_text is not None:
            header_paragraph = _prepare_paragraph(
                section.header, clear_existing=clear_existing
            )
            header_paragraph.text = header_text
            header_paragraph.alignment = normalized_header_alignment

        should_touch_footer = footer_text is not None or include_page_number
        if should_touch_footer:
            footer_paragraph = _prepare_paragraph(
                section.footer, clear_existing=clear_existing
            )
            if footer_text is not None:
                footer_paragraph.text = footer_text
            footer_paragraph.alignment = normalized_footer_alignment
            if include_page_number:
                if footer_paragraph.text:
                    footer_paragraph.add_run(" ")
                _append_page_number_field(footer_paragraph)

        if start_page_number is not None:
            _set_start_page_number(section, start_page_number)

        updated_indices.append(idx)

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "section_count": len(doc.sections),
        "updated_sections": updated_indices,
        "header_updated": header_text is not None,
        "footer_updated": footer_text is not None or include_page_number,
        "include_page_number": include_page_number,
        "start_page_number": start_page_number,
        "different_first_page": different_first_page,
        "different_odd_even": different_odd_even,
    }

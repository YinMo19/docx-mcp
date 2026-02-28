from __future__ import annotations

import re
from collections.abc import Iterable

from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from docx_mcp.errors import DocxMCPError

HEX_COLOR_RE = re.compile(r"^[0-9A-Fa-f]{6}$")


def iter_paragraphs_in_table(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_paragraphs_in_table(nested)


def iter_all_paragraphs(doc) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from iter_paragraphs_in_table(table)


def validate_color(color: str) -> str:
    normalized = color.strip().lstrip("#")
    if not HEX_COLOR_RE.fullmatch(normalized):
        raise DocxMCPError(
            code="INVALID_COLOR",
            message="Color must be a 6-char hex RGB value, e.g. '000000' or '#1A2B3C'.",
            details={"color": color},
        )
    return normalized.upper()


def apply_run_format(
    run,
    *,
    font_name: str | None = None,
    font_size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    if font_name:
        run.font.name = font_name
    if font_size is not None:
        if font_size <= 0:
            raise DocxMCPError(
                code="INVALID_FONT_SIZE",
                message="font_size must be greater than 0.",
                details={"font_size": font_size},
            )
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(validate_color(color))


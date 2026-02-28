from __future__ import annotations

from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from docx_mcp.errors import DocxMCPError
from docx_mcp.services.document_io import open_document, resolve_for_write, save_document
from docx_mcp.services.format_utils import apply_run_format, validate_color

_BORDER_STYLE_MAP = {
    "single": "single",
    "dashed": "dashed",
    "dotted": "dotted",
    "double": "double",
    "none": "nil",
}

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_cell_shading(cell, fill_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    current = tc_pr.find(qn("w:shd"))
    if current is not None:
        tc_pr.remove(current)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def _set_table_borders(table, border_style: str) -> None:
    style = _BORDER_STYLE_MAP.get(border_style.lower())
    if style is None:
        raise DocxMCPError(
            code="INVALID_BORDER_STYLE",
            message=f"Unsupported border_style: {border_style}",
            details={"supported": sorted(_BORDER_STYLE_MAP.keys())},
        )

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    current = tbl_pr.find(qn("w:tblBorders"))
    if current is not None:
        tbl_pr.remove(current)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def format_table(
    filename: str,
    table_index: int,
    border_style: str = "single",
    has_header_row: bool = True,
    shading: list[str] | None = None,
    header_fill_color: str = "D9E2F3",
    header_text_color: str | None = None,
    auto_fit: bool | None = None,
    output_filename: str | None = None,
) -> dict[str, Any]:
    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)

    if table_index < 0 or table_index >= len(doc.tables):
        raise DocxMCPError(
            code="TABLE_INDEX_OUT_OF_RANGE",
            message=f"table_index {table_index} is out of range",
            details={"table_count": len(doc.tables)},
        )

    table = doc.tables[table_index]
    _set_table_borders(table, border_style=border_style)

    if auto_fit is not None:
        table.autofit = bool(auto_fit)

    header_color = validate_color(header_fill_color)
    header_font_color = validate_color(header_text_color) if header_text_color else None

    shaded_rows = 0
    if has_header_row and table.rows:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, header_color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    if header_font_color:
                        run.font.color.rgb = RGBColor.from_string(header_font_color)
        shaded_rows += 1

    alternating = None
    if shading:
        if len(shading) < 2:
            raise DocxMCPError(
                code="INVALID_SHADING",
                message="shading must include at least 2 colors.",
                details={"shading": shading},
            )
        alternating = [validate_color(shading[0]), validate_color(shading[1])]

    if alternating:
        start_idx = 1 if has_header_row else 0
        for idx, row in enumerate(table.rows[start_idx:], start=start_idx):
            fill = alternating[(idx - start_idx) % 2]
            for cell in row.cells:
                _set_cell_shading(cell, fill)
            shaded_rows += 1

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "table_index": table_index,
        "rows": len(table.rows),
        "cols": max((len(r.cells) for r in table.rows), default=0),
        "border_style": border_style,
        "has_header_row": has_header_row,
        "shaded_rows": shaded_rows,
    }


def set_paragraph_format(
    filename: str,
    paragraph_indices: list[int] | None = None,
    contains_text: str | None = None,
    font_name: str | None = None,
    font_size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    alignment: str | None = None,
    line_spacing: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
    left_indent_pt: float | None = None,
    right_indent_pt: float | None = None,
    first_line_indent_pt: float | None = None,
    output_filename: str | None = None,
) -> dict[str, Any]:
    if not paragraph_indices and not contains_text:
        raise DocxMCPError(
            code="MISSING_SELECTOR",
            message="Provide paragraph_indices or contains_text.",
        )

    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)

    index_set = set(paragraph_indices or [])
    for idx in index_set:
        if idx < 0 or idx >= len(doc.paragraphs):
            raise DocxMCPError(
                code="PARAGRAPH_INDEX_OUT_OF_RANGE",
                message=f"Paragraph index out of range: {idx}",
                details={"paragraph_count": len(doc.paragraphs)},
            )

    normalized_alignment = None
    if alignment:
        normalized_alignment = _ALIGNMENT_MAP.get(alignment.lower())
        if normalized_alignment is None:
            raise DocxMCPError(
                code="INVALID_ALIGNMENT",
                message=f"Unsupported alignment: {alignment}",
                details={"supported": sorted(_ALIGNMENT_MAP.keys())},
            )

    if line_spacing is not None and line_spacing <= 0:
        raise DocxMCPError(
            code="INVALID_LINE_SPACING",
            message="line_spacing must be greater than 0.",
            details={"line_spacing": line_spacing},
        )

    matched_indices: list[int] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        match_by_index = idx in index_set
        match_by_text = bool(contains_text and contains_text in (paragraph.text or ""))
        if not (match_by_index or match_by_text):
            continue

        matched_indices.append(idx)
        p_format = paragraph.paragraph_format

        if normalized_alignment is not None:
            paragraph.alignment = normalized_alignment
        if line_spacing is not None:
            p_format.line_spacing = line_spacing
        if space_before_pt is not None:
            p_format.space_before = Pt(space_before_pt)
        if space_after_pt is not None:
            p_format.space_after = Pt(space_after_pt)
        if left_indent_pt is not None:
            p_format.left_indent = Pt(left_indent_pt)
        if right_indent_pt is not None:
            p_format.right_indent = Pt(right_indent_pt)
        if first_line_indent_pt is not None:
            p_format.first_line_indent = Pt(first_line_indent_pt)

        for run in paragraph.runs:
            apply_run_format(
                run,
                font_name=font_name,
                font_size=font_size,
                bold=bold,
                italic=italic,
                color=color,
            )

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "matched_count": len(matched_indices),
        "matched_indices": matched_indices,
    }


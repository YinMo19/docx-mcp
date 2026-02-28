from __future__ import annotations

import re
from typing import Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_mcp.errors import DocxMCPError
from docx_mcp.services.document_io import open_document, resolve_for_write, save_document

_BOOKMARK_NAME_RE = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")


def _validate_heading_range(start: int, end: int) -> None:
    if not (1 <= start <= 9 and 1 <= end <= 9 and start <= end):
        raise DocxMCPError(
            code="INVALID_HEADING_RANGE",
            message="Heading range must satisfy 1 <= start <= end <= 9.",
            details={"heading_start": start, "heading_end": end},
        )


def _validate_bookmark_name(bookmark_name: str) -> None:
    if not _BOOKMARK_NAME_RE.fullmatch(bookmark_name):
        raise DocxMCPError(
            code="INVALID_BOOKMARK_NAME",
            message="bookmark_name must match [A-Za-z_][A-Za-z0-9_]*",
            details={"bookmark_name": bookmark_name},
        )


def _append_field_code(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def insert_table_of_contents(
    filename: str,
    *,
    heading_start: int = 1,
    heading_end: int = 3,
    title_text: str | None = "Table of Contents",
    add_page_break_before: bool = False,
    output_filename: str | None = None,
) -> dict[str, Any]:
    _validate_heading_range(heading_start, heading_end)

    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)

    if add_page_break_before:
        doc.add_page_break()
    if title_text:
        doc.add_paragraph(title_text)

    toc_paragraph = doc.add_paragraph()
    _append_field_code(
        toc_paragraph,
        instruction=f'TOC \\o "{heading_start}-{heading_end}" \\h \\z \\u',
    )

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "heading_start": heading_start,
        "heading_end": heading_end,
        "title_text": title_text,
        "toc_paragraph_index": len(doc.paragraphs) - 1,
    }


def add_sequence_caption(
    filename: str,
    *,
    label: str,
    caption_text: str,
    seq_identifier: str | None = None,
    separator: str = ": ",
    output_filename: str | None = None,
) -> dict[str, Any]:
    if not label.strip():
        raise DocxMCPError(
            code="INVALID_LABEL",
            message="label cannot be empty.",
        )
    if not caption_text.strip():
        raise DocxMCPError(
            code="INVALID_CAPTION_TEXT",
            message="caption_text cannot be empty.",
        )

    sequence_name = seq_identifier.strip() if seq_identifier else label.strip()
    if not sequence_name:
        raise DocxMCPError(
            code="INVALID_SEQUENCE",
            message="seq_identifier cannot be empty when provided.",
        )

    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)

    paragraph = doc.add_paragraph()
    try:
        paragraph.style = doc.styles["Caption"]
    except KeyError:
        pass

    paragraph.add_run(f"{label.strip()} ")
    _append_field_code(paragraph, instruction=f"SEQ {sequence_name} \\* ARABIC")
    paragraph.add_run(f"{separator}{caption_text.strip()}")

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "label": label.strip(),
        "seq_identifier": sequence_name,
        "paragraph_index": len(doc.paragraphs) - 1,
    }


def _next_bookmark_id(doc) -> int:
    bookmark_starts = doc._element.xpath(".//w:bookmarkStart")
    max_id = 0
    for el in bookmark_starts:
        raw = el.get(qn("w:id"))
        if raw is None:
            continue
        try:
            max_id = max(max_id, int(raw))
        except ValueError:
            continue
    return max_id + 1


def add_bookmark_to_paragraph(
    filename: str,
    *,
    paragraph_index: int,
    bookmark_name: str,
    output_filename: str | None = None,
) -> dict[str, Any]:
    _validate_bookmark_name(bookmark_name)

    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise DocxMCPError(
            code="PARAGRAPH_INDEX_OUT_OF_RANGE",
            message=f"Paragraph index out of range: {paragraph_index}",
            details={"paragraph_count": len(doc.paragraphs)},
        )

    paragraph = doc.paragraphs[paragraph_index]
    bookmark_id = _next_bookmark_id(doc)

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "paragraph_index": paragraph_index,
        "bookmark_name": bookmark_name,
        "bookmark_id": bookmark_id,
    }


def insert_ref_field(
    filename: str,
    *,
    bookmark_name: str,
    prefix_text: str | None = None,
    hyperlink: bool = True,
    output_filename: str | None = None,
) -> dict[str, Any]:
    _validate_bookmark_name(bookmark_name)

    src_path, dst_path = resolve_for_write(filename, output_filename)
    doc = open_document(src_path)

    paragraph = doc.add_paragraph()
    if prefix_text:
        paragraph.add_run(prefix_text)
        paragraph.add_run(" ")

    instruction = f"REF {bookmark_name}"
    if hyperlink:
        instruction += " \\h"
    _append_field_code(paragraph, instruction=instruction)

    save_document(doc, dst_path)
    return {
        "source_path": str(src_path),
        "output_path": str(dst_path),
        "bookmark_name": bookmark_name,
        "paragraph_index": len(doc.paragraphs) - 1,
        "hyperlink": hyperlink,
    }

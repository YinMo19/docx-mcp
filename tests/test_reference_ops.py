from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document

from docx_mcp.errors import DocxMCPError
from docx_mcp.server import create_server
from docx_mcp.services.reference_ops import (
    add_bookmark_to_paragraph,
    add_sequence_caption,
    insert_ref_field,
    insert_table_of_contents,
)


def _seed_doc(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Body paragraph for references.")
    doc.save(path)
    return path


def _call_tool(name: str, arguments: dict[str, object]) -> dict[str, object]:
    server = create_server()
    _, payload = asyncio.run(server.call_tool(name, arguments))
    assert isinstance(payload, dict)
    return payload


def test_insert_table_of_contents_adds_toc_field(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "toc.docx"

    result = insert_table_of_contents(
        filename=str(source),
        heading_start=1,
        heading_end=3,
        title_text="Contents",
        output_filename=str(output),
    )
    assert result["toc_paragraph_index"] >= 0

    doc = Document(output)
    assert any(p.text == "Contents" for p in doc.paragraphs)
    assert any('TOC \\o "1-3"' in p._p.xml for p in doc.paragraphs)


def test_add_sequence_caption_inserts_seq_field(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "caption.docx"

    result = add_sequence_caption(
        filename=str(source),
        label="Figure",
        caption_text="System architecture",
        output_filename=str(output),
    )
    assert result["label"] == "Figure"

    doc = Document(output)
    paragraph = doc.paragraphs[-1]
    assert paragraph.text.startswith("Figure")
    assert "System architecture" in paragraph.text
    assert "SEQ Figure" in paragraph._p.xml


def test_bookmark_and_ref_field_flow(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    bookmarked = tmp_path / "bookmarked.docx"
    referenced = tmp_path / "referenced.docx"

    add_bookmark_to_paragraph(
        filename=str(source),
        paragraph_index=1,
        bookmark_name="target_para",
        output_filename=str(bookmarked),
    )
    insert_ref_field(
        filename=str(bookmarked),
        bookmark_name="target_para",
        prefix_text="See",
        output_filename=str(referenced),
    )

    doc = Document(referenced)
    body_xml = doc._element.xml
    assert 'w:name="target_para"' in body_xml
    assert "REF target_para \\h" in body_xml


def test_insert_table_of_contents_rejects_invalid_range(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    with pytest.raises(DocxMCPError) as exc:
        insert_table_of_contents(filename=str(source), heading_start=3, heading_end=1)
    assert exc.value.code == "INVALID_HEADING_RANGE"


def test_reference_tools_accept_string_inputs(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    step1 = tmp_path / "step1.docx"
    step2 = tmp_path / "step2.docx"
    step3 = tmp_path / "step3.docx"

    toc_payload = _call_tool(
        "insert_table_of_contents",
        {
            "filename": str(source),
            "heading_start": "1",
            "heading_end": "3",
            "add_page_break_before": "false",
            "output_filename": str(step1),
        },
    )
    assert toc_payload["ok"] is True

    bookmark_payload = _call_tool(
        "add_bookmark_to_paragraph",
        {
            "filename": str(step1),
            "paragraph_index": "1",
            "bookmark_name": "chapter_1",
            "output_filename": str(step2),
        },
    )
    assert bookmark_payload["ok"] is True

    ref_payload = _call_tool(
        "insert_ref_field",
        {
            "filename": str(step2),
            "bookmark_name": "chapter_1",
            "prefix_text": "See",
            "hyperlink": "true",
            "output_filename": str(step3),
        },
    )
    assert ref_payload["ok"] is True

    doc = Document(step3)
    assert any("REF chapter_1 \\h" in p._p.xml for p in doc.paragraphs)

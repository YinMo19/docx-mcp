from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_mcp.services.read_ops import (
    find_text_in_document,
    get_document_info,
    get_document_outline,
    get_document_text,
    list_available_documents,
)


def _build_sample_docx(path: Path) -> Path:
    doc = Document()
    doc.core_properties.author = "tester"
    doc.core_properties.title = "Sample Doc"

    doc.add_heading("Repository Guidelines", level=1)
    doc.add_paragraph("This project parses DOCX files.")
    doc.add_paragraph("CaseSensitiveToken")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "含义"
    table.cell(1, 0).text = "username"
    table.cell(1, 1).text = "账号"

    doc.save(path)
    return path


def test_list_available_documents_filters_docx(tmp_path: Path) -> None:
    sample = _build_sample_docx(tmp_path / "sample.docx")
    (tmp_path / "notes.txt").write_text("ignore me")

    result = list_available_documents(str(tmp_path))
    names = [doc["name"] for doc in result["documents"]]

    assert result["count"] == 1
    assert names == [sample.name]


def test_get_document_info_reads_metadata(tmp_path: Path) -> None:
    sample = _build_sample_docx(tmp_path / "sample.docx")

    info = get_document_info(str(sample))
    assert info["author"] == "tester"
    assert info["title"] == "Sample Doc"
    assert info["table_count"] == 1
    assert info["paragraph_count"] >= 3


def test_get_document_text_includes_table_cells(tmp_path: Path) -> None:
    sample = _build_sample_docx(tmp_path / "sample.docx")

    text_result = get_document_text(str(sample))
    text = text_result["text"]

    assert "Repository Guidelines" in text
    assert "This project parses DOCX files." in text
    assert "username" in text


def test_get_document_outline_contains_paragraphs_and_tables(tmp_path: Path) -> None:
    sample = _build_sample_docx(tmp_path / "sample.docx")
    outline = get_document_outline(str(sample))

    assert len(outline["paragraphs"]) >= 3
    assert len(outline["tables"]) == 1
    assert outline["tables"][0]["rows"] == 2
    assert outline["tables"][0]["columns"] == 2


def test_find_text_in_document_case_and_whole_word(tmp_path: Path) -> None:
    sample = _build_sample_docx(tmp_path / "sample.docx")

    case_insensitive = find_text_in_document(
        filename=str(sample),
        text_to_find="casesensitivetoken",
        match_case=False,
    )
    assert case_insensitive["total_count"] == 1

    case_sensitive_fail = find_text_in_document(
        filename=str(sample),
        text_to_find="casesensitivetoken",
        match_case=True,
    )
    assert case_sensitive_fail["total_count"] == 0

    whole_word = find_text_in_document(
        filename=str(sample),
        text_to_find="user",
        whole_word=True,
    )
    assert whole_word["total_count"] == 0

    partial = find_text_in_document(
        filename=str(sample),
        text_to_find="user",
        whole_word=False,
    )
    assert partial["total_count"] >= 1


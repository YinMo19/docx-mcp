from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from docx_mcp.errors import DocxMCPError
from docx_mcp.server import create_server
from docx_mcp.services.section_ops import set_headers_footers


def _seed_sections_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Section 0 body")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Section 1 body")
    doc.save(path)
    return path


def _call_tool(name: str, arguments: dict[str, object]) -> dict[str, object]:
    server = create_server()
    _, payload = asyncio.run(server.call_tool(name, arguments))
    assert isinstance(payload, dict)
    return payload


def test_set_headers_footers_updates_selected_section(tmp_path: Path) -> None:
    source = _seed_sections_doc(tmp_path / "source.docx")
    output = tmp_path / "section_update.docx"

    result = set_headers_footers(
        filename=str(source),
        section_indices=[1],
        header_text="Header S1",
        footer_text="Footer S1",
        include_page_number=True,
        start_page_number=5,
        unlink_from_previous=True,
        output_filename=str(output),
    )
    assert result["updated_sections"] == [1]

    doc = Document(output)
    section0 = doc.sections[0]
    section1 = doc.sections[1]

    assert section0.header.paragraphs[0].text == ""
    assert section1.header.paragraphs[0].text == "Header S1"
    assert section1.footer.paragraphs[0].text.startswith("Footer S1")
    assert "PAGE" in section1.footer.paragraphs[0]._p.xml

    pg_num_type = section1._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type is not None
    assert pg_num_type.get(qn("w:start")) == "5"


def test_set_headers_footers_rejects_out_of_range_index(tmp_path: Path) -> None:
    source = _seed_sections_doc(tmp_path / "source.docx")

    with pytest.raises(DocxMCPError) as exc:
        set_headers_footers(filename=str(source), section_indices=[99], header_text="x")
    assert exc.value.code == "SECTION_INDEX_OUT_OF_RANGE"


def test_set_headers_footers_tool_accepts_string_inputs(tmp_path: Path) -> None:
    source = _seed_sections_doc(tmp_path / "source.docx")
    output = tmp_path / "section_tool.docx"

    payload = _call_tool(
        "set_headers_footers",
        {
            "filename": str(source),
            "section_indices": "1",
            "header_text": "Tool Header",
            "footer_text": "Tool Footer",
            "include_page_number": "true",
            "start_page_number": "7",
            "different_first_page": "false",
            "different_odd_even": "true",
            "unlink_from_previous": "true",
            "output_filename": str(output),
        },
    )

    assert payload["ok"] is True
    doc = Document(output)
    section1 = doc.sections[1]
    assert section1.header.paragraphs[0].text == "Tool Header"
    assert "PAGE" in section1.footer.paragraphs[0]._p.xml
    assert doc.settings.odd_and_even_pages_header_footer is True

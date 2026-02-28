from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_mcp.errors import DocxMCPError
from docx_mcp.services.style_ops import format_table, set_paragraph_format


def _seed_style_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Intro paragraph")
    doc.add_paragraph("Target paragraph")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "R1C1"
    table.cell(1, 1).text = "R1C2"
    table.cell(2, 0).text = "R2C1"
    table.cell(2, 1).text = "R2C2"
    doc.save(path)
    return path


def _cell_fill(cell) -> str | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def test_set_paragraph_format_applies_run_and_paragraph_properties(tmp_path: Path) -> None:
    source = _seed_style_doc(tmp_path / "source.docx")
    output = tmp_path / "formatted_para.docx"

    result = set_paragraph_format(
        filename=str(source),
        contains_text="Target",
        font_name="Times New Roman",
        font_size=14,
        bold=True,
        color="445566",
        alignment="center",
        line_spacing=1.5,
        space_before_pt=6,
        space_after_pt=8,
        first_line_indent_pt=12,
        output_filename=str(output),
    )

    assert result["matched_count"] == 1
    doc = Document(output)
    target = doc.paragraphs[1]
    run = target.runs[0]

    assert run.font.name == "Times New Roman"
    assert int(run.font.size.pt) == 14
    assert run.bold is True
    assert str(run.font.color.rgb) == "445566"
    assert target.paragraph_format.line_spacing == 1.5
    assert int(target.paragraph_format.space_before.pt) == 6
    assert int(target.paragraph_format.space_after.pt) == 8
    assert int(target.paragraph_format.first_line_indent.pt) == 12


def test_format_table_applies_borders_and_shading(tmp_path: Path) -> None:
    source = _seed_style_doc(tmp_path / "source.docx")
    output = tmp_path / "formatted_table.docx"

    result = format_table(
        filename=str(source),
        table_index=0,
        border_style="single",
        has_header_row=True,
        shading=["FFFFFF", "EEEEEE"],
        header_fill_color="AABBCC",
        header_text_color="112233",
        output_filename=str(output),
    )

    assert result["table_index"] == 0
    doc = Document(output)
    table = doc.tables[0]

    assert _cell_fill(table.cell(0, 0)) == "AABBCC"
    assert _cell_fill(table.cell(1, 0)) == "FFFFFF"
    assert _cell_fill(table.cell(2, 0)) == "EEEEEE"

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    top = borders.find(qn("w:top")) if borders is not None else None
    assert top is not None
    assert top.get(qn("w:val")) == "single"


def test_set_paragraph_format_requires_selector(tmp_path: Path) -> None:
    source = _seed_style_doc(tmp_path / "source.docx")

    with pytest.raises(DocxMCPError) as exc:
        set_paragraph_format(filename=str(source))

    assert exc.value.code == "MISSING_SELECTOR"


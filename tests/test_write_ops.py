from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_mcp.services.write_ops import (
    add_heading,
    add_paragraph,
    add_table,
    search_and_replace,
)


def _seed_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("foo in paragraph")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "foo in cell"
    doc.save(path)
    return path


def test_search_and_replace_updates_paragraphs_and_cells(tmp_path: Path) -> None:
    source = _seed_docx(tmp_path / "source.docx")
    output = tmp_path / "replaced.docx"

    result = search_and_replace(
        filename=str(source),
        find_text="foo",
        replace_text="bar",
        output_filename=str(output),
    )

    assert result["replacement_count"] >= 2
    updated = Document(output)
    assert "bar in paragraph" in updated.paragraphs[0].text
    assert updated.tables[0].cell(0, 0).text == "bar in cell"


def test_add_paragraph_appends_with_format(tmp_path: Path) -> None:
    source = _seed_docx(tmp_path / "source.docx")
    output = tmp_path / "paragraph.docx"

    add_paragraph(
        filename=str(source),
        text="new paragraph",
        font_name="Times New Roman",
        font_size=12,
        bold=True,
        color="112233",
        output_filename=str(output),
    )

    doc = Document(output)
    paragraph = doc.paragraphs[-1]
    run = paragraph.runs[0]
    assert paragraph.text == "new paragraph"
    assert run.bold is True
    assert run.font.name == "Times New Roman"
    assert int(run.font.size.pt) == 12
    assert str(run.font.color.rgb) == "112233"


def test_add_heading_appends_heading_level(tmp_path: Path) -> None:
    source = _seed_docx(tmp_path / "source.docx")
    output = tmp_path / "heading.docx"

    add_heading(
        filename=str(source),
        text="Section 1",
        level=2,
        output_filename=str(output),
    )

    doc = Document(output)
    heading = doc.paragraphs[-1]
    assert heading.text == "Section 1"
    assert heading.style.name.lower().startswith("heading")


def test_add_table_appends_and_fills_data(tmp_path: Path) -> None:
    source = _seed_docx(tmp_path / "source.docx")
    output = tmp_path / "table.docx"

    add_table(
        filename=str(source),
        rows=2,
        cols=2,
        data=[["A", "B"], ["C", "D"]],
        output_filename=str(output),
    )

    doc = Document(output)
    table = doc.tables[-1]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.cell(0, 0).text == "A"
    assert table.cell(1, 1).text == "D"


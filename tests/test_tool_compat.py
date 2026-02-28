from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_mcp.errors import DocxMCPError
from docx_mcp.server import create_server
from docx_mcp.tools.compat import (
    parse_auto_fit,
    parse_bool,
    parse_matrix,
    parse_optional_bool,
    parse_optional_float,
    parse_paragraph_indices,
)


def _seed_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Intro")
    doc.add_paragraph("Target text")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "R1C1"
    table.cell(1, 1).text = "R1C2"
    table.cell(2, 0).text = "R2C1"
    table.cell(2, 1).text = "R2C2"
    doc.save(path)
    return path


def _call_tool(name: str, arguments: dict[str, object]) -> dict[str, object]:
    server = create_server()
    _, payload = asyncio.run(server.call_tool(name, arguments))
    return payload


def _cell_fill(cell) -> str | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def test_compat_parsers_handle_legacy_string_inputs() -> None:
    assert parse_bool("yes", field="flag") is True
    assert parse_optional_bool("", field="flag") is None
    assert parse_optional_float(" 2.5 ", field="size") == 2.5
    assert parse_paragraph_indices("1, 2;3") == [1, 2, 3]
    assert parse_matrix("A,B\nC,D", field="data") == [["A", "B"], ["C", "D"]]
    assert parse_auto_fit("fixed") is False


def test_compat_paragraph_indices_invalid_raises_structured_error() -> None:
    with pytest.raises(DocxMCPError) as exc:
        parse_paragraph_indices("1, x")
    assert exc.value.code == "INVALID_PARAGRAPH_INDICES"


def test_add_table_tool_accepts_string_dimensions_and_matrix(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "table.docx"

    payload = _call_tool(
        "add_table",
        {
            "filename": str(source),
            "rows": "2",
            "cols": "2",
            "data": '[["A","B"],["C","D"]]',
            "output_filename": str(output),
        },
    )

    assert payload["ok"] is True
    doc = Document(output)
    table = doc.tables[-1]
    assert table.cell(0, 0).text == "A"
    assert table.cell(1, 1).text == "D"


def test_set_paragraph_format_tool_accepts_string_style_values(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "paragraph.docx"

    payload = _call_tool(
        "set_paragraph_format",
        {
            "filename": str(source),
            "paragraph_indices": "1",
            "font_name": "Times New Roman",
            "font_size": "15",
            "bold": "true",
            "line_spacing": "1.25",
            "space_after_pt": "6",
            "output_filename": str(output),
        },
    )

    assert payload["ok"] is True
    doc = Document(output)
    paragraph = doc.paragraphs[1]
    run = paragraph.runs[0]
    assert run.bold is True
    assert run.font.name == "Times New Roman"
    assert int(run.font.size.pt) == 15
    assert paragraph.paragraph_format.line_spacing == 1.25
    assert int(paragraph.paragraph_format.space_after.pt) == 6


def test_format_table_tool_accepts_string_boolean_and_shading(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "table_style.docx"

    payload = _call_tool(
        "format_table",
        {
            "filename": str(source),
            "table_index": "0",
            "has_header_row": "true",
            "shading": "FFFFFF,EEEEEE",
            "auto_fit": "fixed",
            "output_filename": str(output),
        },
    )

    assert payload["ok"] is True
    doc = Document(output)
    table = doc.tables[0]
    assert _cell_fill(table.cell(1, 0)) == "FFFFFF"
    assert _cell_fill(table.cell(2, 0)) == "EEEEEE"


def test_tool_returns_structured_error_for_invalid_integer(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")

    payload = _call_tool(
        "add_heading",
        {
            "filename": str(source),
            "text": "Heading",
            "level": "not_int",
        },
    )

    assert payload["ok"] is False
    error = payload["error"]
    assert isinstance(error, dict)
    assert error["code"] == "INVALID_INTEGER"

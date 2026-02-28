from __future__ import annotations

import asyncio
from pathlib import Path

import pytest
from docx import Document

from docx_mcp.errors import DocxMCPError
from docx_mcp.server import create_server
from docx_mcp.services.document_style_ops import apply_document_style


def _seed_doc(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Body paragraph one.")
    doc.add_paragraph("Body paragraph two.")
    doc.save(path)
    return path


def _call_tool(name: str, arguments: dict[str, object]) -> dict[str, object]:
    server = create_server()
    _, payload = asyncio.run(server.call_tool(name, arguments))
    assert isinstance(payload, dict)
    return payload


def test_apply_document_style_updates_layout_and_paragraphs(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "styled.docx"

    result = apply_document_style(
        filename=str(source),
        page_size="A4",
        margin_top_cm=2.6,
        margin_bottom_cm=2.2,
        margin_left_cm=3.0,
        margin_right_cm=2.4,
        normal_font_name="宋体",
        normal_western_font_name="Times New Roman",
        normal_font_size_pt=12.0,
        normal_line_spacing=1.5,
        normal_first_line_indent_pt=24.0,
        normal_alignment="justify",
        heading_font_name="黑体",
        heading_western_font_name="Times New Roman",
        heading_1_size_pt=16.0,
        heading_2_size_pt=14.0,
        heading_3_size_pt=12.0,
        output_filename=str(output),
    )

    assert result["page_size"] == "A4"
    assert result["heading_paragraph_count"] == 1
    assert result["body_paragraph_count"] == 2

    doc = Document(output)
    section = doc.sections[0]
    assert section.page_width.mm == pytest.approx(210.0, abs=0.1)
    assert section.page_height.mm == pytest.approx(297.0, abs=0.1)
    assert section.left_margin.cm == pytest.approx(3.0, abs=0.05)
    assert section.right_margin.cm == pytest.approx(2.4, abs=0.05)

    heading_run = doc.paragraphs[0].runs[0]
    assert heading_run.bold is True
    assert heading_run.font.size is not None
    assert int(heading_run.font.size.pt) == 16

    body_paragraph = doc.paragraphs[1]
    body_run = body_paragraph.runs[0]
    assert body_paragraph.paragraph_format.line_spacing == 1.5
    assert body_paragraph.paragraph_format.first_line_indent is not None
    assert int(body_paragraph.paragraph_format.first_line_indent.pt) == 24
    assert body_run.font.name == "Times New Roman"
    assert body_run.font.size is not None
    assert int(body_run.font.size.pt) == 12


def test_apply_document_style_rejects_invalid_page_size(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    with pytest.raises(DocxMCPError) as exc:
        apply_document_style(filename=str(source), page_size="B5")
    assert exc.value.code == "INVALID_PAGE_SIZE"


def test_apply_document_style_tool_accepts_string_inputs(tmp_path: Path) -> None:
    source = _seed_doc(tmp_path / "source.docx")
    output = tmp_path / "tool_styled.docx"

    payload = _call_tool(
        "apply_document_style",
        {
            "filename": str(source),
            "margin_left_cm": "3.2",
            "margin_right_cm": "2.3",
            "normal_font_size_pt": "12",
            "normal_line_spacing": "1.6",
            "normal_first_line_indent_pt": "26",
            "max_heading_level": "3",
            "apply_to_existing_paragraphs": "true",
            "output_filename": str(output),
        },
    )
    assert payload["ok"] is True

    doc = Document(output)
    section = doc.sections[0]
    assert section.left_margin.cm == pytest.approx(3.2, abs=0.05)
    assert section.right_margin.cm == pytest.approx(2.3, abs=0.05)

    body_paragraph = doc.paragraphs[1]
    assert body_paragraph.paragraph_format.line_spacing == 1.6
    assert body_paragraph.paragraph_format.first_line_indent is not None
    assert int(body_paragraph.paragraph_format.first_line_indent.pt) == 26
